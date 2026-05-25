"""docx 文档解析器 — 原始元素提取"""

import tempfile
from enum import Enum
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from pydantic import BaseModel

from src.core.logger import get_logger

logger = get_logger(__name__)


class ElementType(str, Enum):
    PARAGRAPH = "paragraph"
    TABLE = "table"
    IMAGE = "image"


class RawDocumentElement(BaseModel):
    """原始文档元素，parser 的统一输出"""

    type: ElementType
    content: str
    image_ref: str | None = None
    index: int
    page: int | None = None
    metadata: dict = {}


class DocxParser:
    """从 docx 文件中按顺序提取段落、表格、图片"""

    def __init__(self, image_dir: str | Path | None = None):
        self.image_dir = Path(image_dir) if image_dir else Path(tempfile.mkdtemp(prefix="docx_images_"))
        self.image_dir.mkdir(parents=True, exist_ok=True)

    def parse(self, file_path: str | Path) -> list[RawDocumentElement]:
        doc = Document(str(file_path))
        elements: list[RawDocumentElement] = []
        index = 0

        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                # 段落：可能包含文本和/或图片
                paragraph = self._find_paragraph(doc, child)
                if paragraph is None:
                    continue

                images = self._extract_images_from_paragraph(doc, child)
                text = paragraph.text.strip()

                if images:
                    for img_ref in images:
                        elements.append(
                            RawDocumentElement(
                                type=ElementType.IMAGE,
                                content=text or "",
                                image_ref=img_ref,
                                index=index,
                                metadata=self._paragraph_metadata(paragraph),
                            )
                        )
                        index += 1
                elif text:
                    elements.append(
                        RawDocumentElement(
                            type=ElementType.PARAGRAPH,
                            content=text,
                            index=index,
                            metadata=self._paragraph_metadata(paragraph),
                        )
                    )
                    index += 1

            elif tag == "tbl":
                table = self._find_table(doc, child)
                if table:
                    md = self._table_to_markdown(table)
                    if md.strip():
                        elements.append(
                            RawDocumentElement(
                                type=ElementType.TABLE,
                                content=md,
                                index=index,
                                metadata={"row_count": len(table.rows), "col_count": len(table.columns)},
                            )
                        )
                        index += 1

        logger.info("解析完成: %s，共 %d 个元素", file_path, len(elements))
        return elements

    # ── 内部方法 ──

    def _find_paragraph(self, doc: Document, p_element):
        """通过 XML 元素找到对应的 Paragraph 对象"""
        for p in doc.paragraphs:
            if p._element is p_element:
                return p
        return None

    def _find_table(self, doc: Document, tbl_element):
        """通过 XML 元素找到对应的 Table 对象"""
        for t in doc.tables:
            if t._element is tbl_element:
                return t
        return None

    def _extract_images_from_paragraph(self, doc: Document, p_element) -> list[str]:
        """从段落中提取所有图片，保存到临时目录，返回文件路径列表"""
        refs: list[str] = []
        blips = p_element.findall(f".//{qn('a:blip')}")
        for i, blip in enumerate(blips):
            embed_id = blip.get(qn("r:embed"))
            if not embed_id:
                continue
            try:
                image_part = doc.part.related_parts[embed_id]
                img_bytes = image_part.blob
                ext = Path(image_part.partname).suffix or ".png"
                img_name = f"{id(p_element)}_{i}{ext}"
                img_path = self.image_dir / img_name
                img_path.write_bytes(img_bytes)
                refs.append(str(img_path))
            except Exception:
                logger.warning("图片提取失败: embed_id=%s", embed_id)
        return refs

    @staticmethod
    def _table_to_markdown(table) -> str:
        """将表格转换为 markdown 格式"""
        rows_data: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return ""

        lines: list[str] = []
        # 表头
        lines.append("| " + " | ".join(rows_data[0]) + " |")
        lines.append("| " + " | ".join("---" for _ in rows_data[0]) + " |")
        # 数据行
        for row in rows_data[1:]:
            # 补齐列数
            while len(row) < len(rows_data[0]):
                row.append("")
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)

    @staticmethod
    def _paragraph_metadata(paragraph) -> dict:
        meta: dict = {}
        if paragraph.style and paragraph.style.name:
            meta["style"] = paragraph.style.name
        if paragraph.alignment is not None:
            meta["alignment"] = str(paragraph.alignment)
        return meta
